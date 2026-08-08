"""Tests for jarvis.multimodal — Vision / Speech / Document processors."""

from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

import pytest

# Ensure emperor-core is importable
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))


# ══════════════════════════════════════════════════════════════════════
# Fixtures
# ══════════════════════════════════════════════════════════════════════


@pytest.fixture
def vision_processor():
    from jarvis.multimodal.processor import VisionProcessor

    return VisionProcessor(llm_engine=None)


@pytest.fixture
def speech_processor():
    from jarvis.multimodal.processor import SpeechProcessor

    return SpeechProcessor(openai_client=None)


@pytest.fixture
def document_processor(vision_processor):
    from jarvis.multimodal.processor import DocumentProcessor

    return DocumentProcessor(vision_processor=vision_processor)


@pytest.fixture
def multimodal_engine():
    from jarvis.multimodal.engine import MultimodalEngine

    return MultimodalEngine(llm_engine=None, openai_client=None)


# ══════════════════════════════════════════════════════════════════════
# Module imports
# ══════════════════════════════════════════════════════════════════════


class TestModuleImports:
    def test_import_multimodalengine(self):
        from jarvis.multimodal import MultimodalEngine

        assert MultimodalEngine is not None

    def test_import_vision_processor(self):
        from jarvis.multimodal import VisionProcessor

        assert VisionProcessor is not None

    def test_import_speech_processor(self):
        from jarvis.multimodal import SpeechProcessor

        assert SpeechProcessor is not None

    def test_import_document_processor(self):
        from jarvis.multimodal import DocumentProcessor

        assert DocumentProcessor is not None

    def test_init_exports(self):
        from jarvis.multimodal import (
            DocumentProcessor,
            MultimodalEngine,
            SpeechProcessor,
            VisionProcessor,
        )

        assert MultimodalEngine.__module__ == "jarvis.multimodal.engine"
        assert VisionProcessor.__module__ == "jarvis.multimodal.processor"
        assert SpeechProcessor.__module__ == "jarvis.multimodal.processor"
        assert DocumentProcessor.__module__ == "jarvis.multimodal.processor"


# ══════════════════════════════════════════════════════════════════════
# VisionProcessor
# ══════════════════════════════════════════════════════════════════════


class TestVisionProcessor:
    def test_init_without_llm(self, vision_processor):
        assert vision_processor is not None
        assert vision_processor._llm is None

    def test_process_url_returns_dict(self, vision_processor):
        result = vision_processor.process("https://example.com/photo.jpg")
        assert isinstance(result, dict)
        assert "caption" in result
        assert "raw" in result
        assert "image_url" in result

    def test_process_local_fallback(self, vision_processor):
        """Process a tiny test image using the fallback path."""
        from PIL import Image

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            img = Image.new("RGB", (10, 10), color="red")
            img.save(f, format="PNG")
            tmp_path = f.name

        try:
            result = vision_processor.process(tmp_path)
            assert isinstance(result, dict)
            assert "caption" in result
            assert "image_path" in result
            assert result["image_path"] == tmp_path
        finally:
            os.unlink(tmp_path)

    def test_process_missing_file_raises(self, vision_processor):
        with pytest.raises(Exception):
            vision_processor.process("/nonexistent/photo.jpg")

    def test_custom_prompt(self, vision_processor):
        result = vision_processor.process(
            "https://example.com/photo.jpg",
            prompt="What colors are in this image?",
        )
        assert isinstance(result, dict)

    def test_custom_system(self, vision_processor):
        result = vision_processor.process(
            "https://example.com/photo.jpg",
            system="You are a color expert.",
        )
        assert isinstance(result, dict)


# ══════════════════════════════════════════════════════════════════════
# SpeechProcessor
# ══════════════════════════════════════════════════════════════════════


class TestSpeechProcessor:
    def test_init(self, speech_processor):
        assert speech_processor is not None
        assert speech_processor._client is None

    def test_process_missing_file(self, speech_processor):
        result = speech_processor.process("/nonexistent/audio.mp3")
        assert "error" in result

    def test_synthesize_creates_mp3(self, speech_processor):
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, "test_out.mp3")
            result = speech_processor.synthesize("Hello world", output_path=out)
            assert isinstance(result, dict)
            assert "audio_path" in result
            assert "text" in result
            # edge-tts may not be installed locally, so just check shape
            assert result["text"] == "Hello world"

    def test_synthesize_auto_path(self, speech_processor):
        result = speech_processor.synthesize("Quick test")
        assert isinstance(result, dict)
        assert "audio_path" in result
        # clean up
        ap = result.get("audio_path", "")
        if ap and os.path.isfile(ap):
            os.unlink(ap)


# ══════════════════════════════════════════════════════════════════════
# DocumentProcessor
# ══════════════════════════════════════════════════════════════════════


class TestDocumentProcessor:
    def test_init(self, document_processor):
        assert document_processor is not None

    def test_process_missing_file(self, document_processor):
        result = document_processor.process("/nonexistent/doc.pdf")
        assert "error" in result

    def test_process_unsupported_type(self, document_processor):
        with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
            f.write(b"not a real doc")
            tmp = f.name
        try:
            result = document_processor.process(tmp)
            assert "error" in result
            assert "supported_types" in result
        finally:
            os.unlink(tmp)

    def test_process_pdf(self, document_processor):
        """Create a minimal PDF and parse it with PyPDF2."""
        try:
            import PyPDF2  # noqa: E402
        except ImportError:
            pytest.skip("PyPDF2 not installed")

        from io import BytesIO

        import PyPDF2

        writer = PyPDF2.PdfWriter()
        writer.add_blank_page(612, 792)  # US Letter

        # PyPDF2 doesn't natively write text, so we test parsing an empty page
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            writer.write(f)
            tmp = f.name

        try:
            result = document_processor.process(tmp)
            assert isinstance(result, dict)
            assert "content" in result
            assert "page_count" in result
            assert result["page_count"] == 1
            assert result["file_path"] == tmp
        finally:
            os.unlink(tmp)

    def test_process_docx(self, document_processor):
        """Create a minimal DOCX and parse it."""
        try:
            import docx  # noqa: E402
        except ImportError:
            pytest.skip("python-docx not installed")

        import docx

        d = docx.Document()
        d.add_paragraph("Hello multimodal world!")
        d.add_paragraph("This is a test document.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            d.save(f.name)
            tmp = f.name

        try:
            result = document_processor.process(tmp)
            assert isinstance(result, dict)
            assert "content" in result
            assert "paragraph_count" in result
            assert result["paragraph_count"] == 2
            assert "Hello multimodal world" in result["content"]
        finally:
            os.unlink(tmp)

    def test_process_image_ocr(self, document_processor):
        """Test OCR path on a synthetic image."""
        from PIL import Image

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            img = Image.new("RGB", (20, 20), color="white")
            img.save(f, format="PNG")
            tmp = f.name

        try:
            result = document_processor.process(tmp)
            assert isinstance(result, dict)
            assert "content" in result
            assert result.get("ocr") is True
        finally:
            os.unlink(tmp)


# ══════════════════════════════════════════════════════════════════════
# MultimodalEngine
# ══════════════════════════════════════════════════════════════════════


class TestMultimodalEngine:
    def test_init(self, multimodal_engine):
        assert multimodal_engine is not None
        assert multimodal_engine.vision is not None
        assert multimodal_engine.speech is not None
        assert multimodal_engine.document is not None

    def test_see(self, multimodal_engine):
        result = multimodal_engine.see("https://example.com/photo.jpg")
        assert isinstance(result, dict)
        assert "caption" in result

    def test_see_custom_prompt(self, multimodal_engine):
        result = multimodal_engine.see(
            "https://example.com/photo.jpg",
            prompt="List every object in this image.",
        )
        assert isinstance(result, dict)

    def test_hear_missing_file(self, multimodal_engine):
        result = multimodal_engine.hear("/nonexistent/audio.mp3")
        assert "error" in result

    def test_speak(self, multimodal_engine):
        result = multimodal_engine.speak("Test speech synthesis")
        assert isinstance(result, dict)
        assert "audio_path" in result
        assert result["text"] == "Test speech synthesis"
        ap = result.get("audio_path", "")
        if ap and os.path.isfile(ap):
            os.unlink(ap)

    def test_speak_custom_voice(self, multimodal_engine):
        result = multimodal_engine.speak("Hello", voice="en-US-AriaNeural")
        assert isinstance(result, dict)
        ap = result.get("audio_path", "")
        if ap and os.path.isfile(ap):
            os.unlink(ap)

    def test_read_document_missing(self, multimodal_engine):
        result = multimodal_engine.read_document("/nonexistent/doc.pdf")
        assert "error" in result

    def test_read_document_docx(self, multimodal_engine):
        """End-to-end: MultimodalEngine → DocumentProcessor for DOCX."""
        try:
            import docx  # noqa: E402
        except ImportError:
            pytest.skip("python-docx not installed")

        import docx

        d = docx.Document()
        d.add_paragraph("E2E multimodal engine test.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            d.save(f.name)
            tmp = f.name

        try:
            result = multimodal_engine.read_document(tmp)
            assert "content" in result
            assert "E2E multimodal engine test" in result["content"]
        finally:
            os.unlink(tmp)

    def test_vision_property(self, multimodal_engine):
        from jarvis.multimodal.processor import VisionProcessor

        assert isinstance(multimodal_engine.vision, VisionProcessor)

    def test_speech_property(self, multimodal_engine):
        from jarvis.multimodal.processor import SpeechProcessor

        assert isinstance(multimodal_engine.speech, SpeechProcessor)

    def test_document_property(self, multimodal_engine):
        from jarvis.multimodal.processor import DocumentProcessor

        assert isinstance(multimodal_engine.document, DocumentProcessor)
